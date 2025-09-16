from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

EXEC_SUMMARY = (
	"DragonShield Antivirus is a modular, AI-powered antivirus with a safe sandbox simulator and optional blockchain-based signature sharing. "
	"This report summarizes the MVP and v2.0 upgrades, architecture, implementation, security posture, and roadmap."
)

MVP_ARCH = """
Figure: MVP Architecture (Mermaid)

graph TD
A[File Input] --> B[Static Scanner]
B -- Known --> V1[Verdict: Malicious]
B -- Unknown --> C[Sandbox Simulator]
C --> D[Behavior Features]
D --> E[AI Classifier]
E --> V2[Verdict]
V1 -->|Opt-in| F[Publish Signature Metadata]
V2 -->|Opt-in| F
F --> G[Mock Blockchain]
""".strip()

V2_ARCH = """
Figure: v2.0 Architecture (Mermaid)

graph LR
subgraph Clients
A1[Native Agent]-- upload -->CAPI
A2[Android Scanner]-- upload -->CAPI
A3[Desktop UI]-- control -->CAPI
end
subgraph Cloud
CAPI[Cloud API]
CFEED[Threat Intel & Model Server]
CSBX[Sandbox Aggregator]
end
subgraph Blockchain
BC1[Staking & Reputation]
BC2[Rewards ERC-20]
end
CAPI-- updates -->A1
CAPI-- updates -->A2
CAPI-- signatures/models -->CFEED
CAPI-- traces -->CSBX
CFEED-- publish metadata -->BC1
""".strip()

SCAN_SEQ = """
Figure: File Scan Sequence (Mermaid)

sequenceDiagram
  participant Client
  participant API
  participant Orchestrator
  participant Static
  participant Sandbox
  participant AI
  Client->>API: POST /scan (file)
  API->>Orchestrator: scan_file
  Orchestrator->>Static: scan
  Static-->>Orchestrator: result
  alt matches
    Orchestrator-->>API: verdict static
  else unknown
    Orchestrator->>Sandbox: run
    Sandbox-->>Orchestrator: behavior
    Orchestrator->>AI: infer
    AI-->>Orchestrator: label
  end
""".strip()

BLOCKCHAIN_FLOW = """
Figure: Blockchain Publish (Mermaid)

sequenceDiagram
  participant Client
  participant API
  participant Chain
  Client->>API: POST /chain/publish
  API->>Chain: publish_signature
  Chain-->>API: entry
  API-->>Client: entry
""".strip()

TECH_TABLE = [
	("Component", "Choice", "Reason"),
	("Core runtime", "Python 3.11", "Rapid prototyping for AI & APIs"),
	("AI", "scikit-learn, optional Torch", "Simple baseline + path to deep models"),
	("Native agent", "Rust", "Performance, portability"),
	("Cloud", "FastAPI + K8s", "Lightweight, scalable"),
	("Sandbox", "Simulator; path to Sysmon/Cuckoo", "Safe dev, production-ready path"),
	("Blockchain", "Solidity (staking, reputation)", "Validator incentives and trust"),
]


def add_heading(doc: Document, text: str, level: int = 1):
	p = doc.add_heading(text, level=level)
	return p


def add_paragraph(doc: Document, text: str):
	return doc.add_paragraph(text)


def add_mermaid_figure(doc: Document, code: str):
	# Render as monospaced text block; viewer can convert to image later
	run = doc.add_paragraph().add_run(code)
	font = run.font
	font.name = 'Courier New'
	font.size = Pt(10)


def add_table(doc: Document, rows: list[tuple[str, str, str]]):
	t = doc.add_table(rows=1, cols=3)
	hdr_cells = t.rows[0].cells
	hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = rows[0]
	for r in rows[1:]:
		cells = t.add_row().cells
		cells[0].text, cells[1].text, cells[2].text = r


def build_docx(path: str):
	doc = Document()
	# Cover page
	title = doc.add_paragraph()
	run = title.add_run("DragonShield Antivirus – Technical Report (MVP to v2.0)\n")
	run.bold = True
	run.font.size = Pt(20)
	sub = doc.add_paragraph("AI-Powered Antivirus with Blockchain Integration")
	sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
	doc.add_paragraph(f"Author: DragonShield Team")
	doc.add_paragraph(f"Date: {date.today().isoformat()}")
	doc.add_page_break()

	# Table of contents note (manual for simplicity)
	add_heading(doc, "Table of Contents", level=1)
	for sec in [
		"Executive Summary","Introduction","System Architecture","Core Modules",
		"Technology Stack","Implementation Overview","Security & Compliance",
		"Roadmap Beyond v2.0","Conclusion"]:
		add_paragraph(doc, f"- {sec}")
	doc.add_page_break()

	# Executive Summary
	add_heading(doc, "Executive Summary", level=1)
	add_paragraph(doc, EXEC_SUMMARY)

	# Introduction
	add_heading(doc, "Introduction", level=1)
	add_paragraph(doc, "Problem: Evolving malware and zero-days require adaptive, trustworthy detection.")
	add_paragraph(doc, "Objectives: Build a safe, modular antivirus with AI behavior analysis and decentralized trust via staking.")
	add_paragraph(doc, "Evolution: MVP (static+sim sandbox+AI+mock chain) -> v2.0 (advanced heuristics, seq AI, cloud, staking).")

	# System Architecture
	add_heading(doc, "System Architecture", level=1)
	add_paragraph(doc, "MVP architecture:")
	add_mermaid_figure(doc, MVP_ARCH)
	add_paragraph(doc, "v2.0 architecture:")
	add_mermaid_figure(doc, V2_ARCH)
	add_paragraph(doc, "Data flows:")
	add_mermaid_figure(doc, SCAN_SEQ)
	add_mermaid_figure(doc, BLOCKCHAIN_FLOW)

	# Core Modules
	add_heading(doc, "Core Modules", level=1)
	add_paragraph(doc, "Static Analysis: signatures + entropy/heuristics (v2).")
	add_paragraph(doc, "Dynamic Sandbox: simulated now; production path via Sysmon/Cuckoo in isolated VMs.")
	add_paragraph(doc, "AI: baseline logistic regression; v2 adds sequence model/LSTM heuristic.")
	add_paragraph(doc, "Cloud Backend: FastAPI services for updates, traces, model distribution.")
	add_paragraph(doc, "Blockchain: staking, reputation, and rewards (Solidity reference).")
	add_paragraph(doc, "UI: web dashboards with behavior timeline; CLI and native agent.")

	# Technology Stack
	add_heading(doc, "Technology Stack", level=1)
	add_table(doc, TECH_TABLE)

	# Implementation Overview
	add_heading(doc, "Implementation Overview", level=1)
	add_paragraph(doc, "Project tree includes endpoint components (Python/Rust), cloud API, blockchain contracts, tests, and deployment manifests.")
	add_paragraph(doc, "Integration: Endpoint performs static+dynamic; unknown samples -> cloud AI; optional publish metadata to blockchain.")

	# Security & Compliance
	add_heading(doc, "Security & Compliance", level=1)
	add_paragraph(doc, "GDPR: data minimization, consent, encryption; anonymized telemetry.")
	add_paragraph(doc, "Encryption: TLS in transit; AES-GCM at rest; KMS-managed keys.")
	add_paragraph(doc, "Sandbox Hardening: VM isolation, RO mounts, eBPF/syscall filters, no outbound network except broker.")
	add_paragraph(doc, "Certifications: AV-TEST, VB100, ISO 27001 preparation.")

	# Roadmap
	add_heading(doc, "Roadmap Beyond v2.0", level=1)
	add_paragraph(doc, "Zero-day exploit detection AI, NIDS, cloud-native workload protection, incident response automation.")

	# Conclusion
	add_heading(doc, "Conclusion", level=1)
	add_paragraph(doc, "DragonShield progressed from MVP to v2.0 with stronger detection, scalability, and decentralized trust foundations.")

	doc.save(path)
	print(f"Report written to {path}")


if __name__ == "__main__":
	build_docx("DragonShield_Technical_Report_v2.docx")
